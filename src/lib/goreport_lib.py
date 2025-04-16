#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Il s'agit de la classe GoReport. GoReport se charge de la connexion au serveur Gophish cible, de la récupération des informations de campagne et de la génération du rapport.
"""

# Standard Libraries
import configparser
import os.path
import sys
from collections import Counter
from datetime import datetime
from types import SimpleNamespace
import json

from src.services.smartphish_service import *
from src.config import Config  # Si config.py est directement sous src/

from src.lib import custom_cve
import time
from docx.enum.text import WD_COLOR_INDEX
from docx import oxml
from docx.opc.constants import RELATIONSHIP_TYPE

# 3rd Party Libraries
import requests
import xlsxwriter
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor

from user_agents import parse

# Fonction utilitaire pour convertir les dictionnaires en objets
def dict_to_obj(d):
    if isinstance(d, dict):
        return SimpleNamespace(**{k: dict_to_obj(v) for k, v in d.items()})
    elif isinstance(d, list):
        return [dict_to_obj(item) for item in d]
    else:
        return d

class Goreport(object):
    """
    This class uses the Gophish library to create a new Gophish API connection
    and queries Gophish for information and results related to the specified
    campaign ID(s).
    """
    # Name of the config file -- default is ``gophish.config``
    goreport_config_file = "gophish.config"
    verbose = False

    # Variables for holding Gophish models
    results = None
    campaign = None
    timeline = None

    # Variables for holding campaign information
    cam_id = None
    cam_url = None
    cam_name = None
    cam_status = None
    launch_date = None
    created_date = None
    cam_page_name = None
    cam_smtp_host = None
    completed_date = None
    cam_redirect_url = None
    cam_from_address = None
    cam_subject_line = None
    cam_template_name = None
    cam_capturing_passwords = None
    cam_capturing_credentials = None

    # Variables and lists for tracking event numbers
    total_sent = 0
    total_opened = 0
    total_targets = 0
    total_clicked = 0
    total_reported = 0
    total_submitted = 0
    total_unique_opened = 0
    total_unique_clicked = 0
    total_unique_reported = 0
    total_unique_submitted = 0
    targets_opened = []
    targets_clicked = []
    targets_reported = []
    targets_submitted = []

    # Lists and dicts for holding prepared report data
    campaign_results_summary = []

    # Lists for holding totals for statistics
    browsers = []
    locations = []
    ip_addresses = []
    ip_and_location = {}
    operating_systems = []

    # Output options
    report_format = None
    output_word_report = None
    output_xlsx_report = None
    xlsx_header_bg_color = "#0085CA"
    xlsx_header_font_color = "#FFFFFF"

    def __init__(self, report_format, config_file, google, verbose):
        """
        Initiate the connection to the Gophish server with the provided host, port,
        and API key and prepare to use the external APIs.
        """
        try:
            # Check if an alternate config file was provided
            if config_file:
                self.goreport_config_file = config_file
            # Open the config file to make sure it exists and is readable
            config = configparser.ConfigParser()
            config.read(self.goreport_config_file)
        except Exception as e:
            print(f"[!] Could not open {self.goreport_config_file} -- make sure it exists and is readable.")
            print(f"L.. Details: {e}")
            sys.exit()

        try:
            # Read in the values from the config file
            GP_HOST = Config.GOPHISH_API_URL
            API_KEY = Config.GOPHISH_API_KEY

        except Exception as e:
            print("[!] There was a problem reading values from the gophish.config file!")
            print(f"L.. Details: {e}")
            sys.exit()

        try:
            # Read in the values from the config file
            self.IPINFO_TOKEN = self.config_section_map(config, 'ipinfo.io')['ipinfo_token']
            if not self.IPINFO_TOKEN:
                self.IPINFO_TOKEN = None
        except Exception as e:
            self.IPINFO_TOKEN = None
            print("[!] No ipinfo.io API token was found in the config. GoReport will not lookup IP addresses with ipinfo.io for additional location data.")
            print(f"L.. Details: {e}")

        try:
            # Read in the values from the config file
            self.GEOLOCATE_TOKEN = self.config_section_map(config, 'Google')['geolocate_key']
            if not self.GEOLOCATE_TOKEN:
                self.GEOLOCATE_TOKEN = None
        except Exception as e:
            self.GEOLOCATE_TOKEN = None
            if google:
                print("[!] No Google Maps API token was found in the config so GoReport will ignore the `--google` flag.")
                print(f"L.. Details: {e}")

        # Set command line options for the GoReport object
        self.google = google
        self.verbose = verbose
        self.report_format = report_format
        # Connect to the Gophish API
        # NOTE: This step succeeds even with a bad API key, so the true test is fetching an ID
        print(f"[+] Connecting to Gophish at {GP_HOST}")
        print(f"L.. The API Authorization endpoint is: {GP_HOST}/api/campaigns/?api_key={API_KEY}")


    def run(self, id_list, combine_reports, set_complete_status):
        """
        Traite les campagnes fournies par leur ID, collecte les informations et génère un rapport.
        Les données sont récupérées via l'API HTTP (module gophish_service) et converties en objets.
        """
        # Affichage des options
        if combine_reports:
            print("[+] Les résultats des campagnes seront combinés dans un seul rapport.")
        if set_complete_status:
            print('[+] Le statut des campagnes sera marqué comme "Complete" après traitement.')

        # Traitement des ID fournis (supporte les plages et les listes séparées par des virgules)
        try:
            temp_id = []
            if "-" in id_list and "," in id_list:
                temp = id_list.split(",")
                for x in temp:
                    if "-" in x:
                        lower, upper = x.split("-")
                        for y in range(int(lower), int(upper) + 1):
                            temp_id.append(str(y))
                    else:
                        temp_id.append(x)
            elif "-" in id_list:
                lower, upper = id_list.split("-")
                for y in range(int(lower), int(upper) + 1):
                    temp_id.append(str(y))
            else:
                temp_id = id_list.split(",")
            id_list = temp_id
        except Exception as e:
            print("[!] Impossible d'interpréter les ID de campagne fournis. Veuillez fournir des ID séparés par des virgules ou des plages (ex. 5,50-55,71).")
            print(f"L.. Détails : {e}")
            sys.exit()

        # Suppression des doublons et tri des ID
        try:
            initial_len = len(id_list)
            id_list = sorted(set(id_list), key=int)
            unique_len = len(id_list)
        except Exception as e:
            temp = []
            for ident in id_list:
                try:
                    int(ident)
                except:
                    temp.append(ident)
            print(f"[!] Il y a {len(temp)} ID de campagne invalides (non entiers).")
            print(f"L.. ID invalides : {','.join(temp)}")
            print(f"L.. Détails : {e}")
            sys.exit()
        print(f"[+] Un total de {initial_len} ID de campagne ont été fournis pour traitement.")
        if initial_len != unique_len:
            dupes = initial_len - unique_len
            print(f"L.. {dupes} ID en double ont été supprimés.")
        print(f"[+] Les campagnes suivantes seront traitées : {','.join(id_list)}")
        if len(id_list) == 1 and combine_reports:
            combine_reports = False

        campaign_counter = 1
        for CAM_ID in id_list:
            print(f"[+] Récupération des résultats pour la campagne ID {CAM_ID} ({campaign_counter}/{len(id_list)}).")
            try:
                # Récupération des détails via l'appel API
                campaign_data = get_campaign(CAM_ID)
                if campaign_data.get("error"):
                    print(f"[!] Échec de récupération des détails pour la campagne ID {CAM_ID}.")
                    print(f"L.. Détails : {campaign_data.get('error')}")
                    if CAM_ID == id_list[-1] and combine_reports:
                        self.generate_report()
                    campaign_counter += 1
                    continue
                else:
                    # Conversion du dictionnaire en objet pour compatibilité avec le reste du code
                    self.campaign = dict_to_obj(campaign_data)
            except Exception as e:
                print(f"[!] Problème lors de la récupération des détails de la campagne ID {CAM_ID} !")
                print(f"L.. Détails : {e}")
                sys.exit()

            # --- Débogage : Affichage et correction de la timeline ---
            try:
                if hasattr(self.campaign, 'timeline') and isinstance(self.campaign.timeline, list):
                    for event in self.campaign.timeline:
                        if isinstance(event.details, str):
                            if event.details.strip() != "":
                                try:
                                    # Convertir la chaîne JSON en dictionnaire
                                    event.details = json.loads(event.details)
                                except Exception as parse_err:
                                    print(f"[DEBUG] Erreur lors du parsing de details pour l'événement {event}: {parse_err}")
                                    event.details = {}
                            else:
                                event.details = {}
            except Exception as e_tl:
                print(f"[DEBUG] Erreur lors du traitement de la timeline: {e_tl}")

            try:
                # Optionnel : vérifier si la campagne récupérée comporte une erreur via un attribut 'success'
                try:
                    if hasattr(self.campaign, 'success') and self.campaign.success is False:
                        print(f"[!] Échec lors de la récupération des résultats pour la campagne ID {CAM_ID}")
                        if hasattr(self.campaign, 'message'):
                            print(f"L.. Détails : {self.campaign.message}")
                        if CAM_ID == id_list[-1] and combine_reports:
                            self.generate_report()
                        campaign_counter += 1
                        continue
                except Exception as e_inner:
                    print(f"[DEBUG] Exception dans le test de l'attribut 'success': {e_inner}")
                    pass

                print("[+] Succès !")
                # Traitement des informations de la campagne
                self.collect_all_campaign_info(combine_reports)
                self.process_timeline_events(combine_reports)
                self.process_results(combine_reports)

                # Marquage de la campagne comme complète si l'option est activée
                if set_complete_status:
                    print(f"[+] Marquage de la campagne ID {CAM_ID} comme terminée.")
                    try:
                        result = complete_campaign(CAM_ID)
                        if not result.get("success", False):
                            print(f"[!] Échec lors du marquage de la campagne ID {CAM_ID} comme terminée.")
                            print(f"L.. Détails : {result.get('message', result.get('content'))}")
                    except Exception as e_mark:
                        print(f"[!] Échec lors du marquage de la campagne ID {CAM_ID} comme terminée.")
                        print(f"L.. Détails : {e_mark}")

                # Génération du rapport
                if CAM_ID == id_list[-1] and combine_reports:
                    self.generate_report()
                elif not combine_reports:
                    self.generate_report()

                campaign_counter += 1
            except Exception as e:
                print(f"[!] Problème lors du traitement de la campagne ID {CAM_ID} !")
                print(f"L.. Détails : {e}")
                sys.exit()


    def lookup_ip(self, ip):
        """Lookup the provided IP address with ipinfo.io for location data.

        Example Result:
            {'ip': '52.44.93.197',
            'hostname': 'ec2-52-44-93-197.compute-1.amazonaws.com',
            'city': 'Beaumont',
            'region': 'Texas',
            'country': 'US',
            'loc': '30.0866,-94.1274',
            'postal': '77702',
            'phone': '409',
            'org': 'AS14618 Amazon.com, Inc.'}
        """
        ipinfo_url = f"https://ipinfo.io/{ip}?token={self.IPINFO_TOKEN}"
        try:
            r = requests.get(ipinfo_url)
            return r.json()
        except Exception as e:
            print(f"[!] Failed to lookup `{ip}` with ipinfo.io.")
            print(f"L.. Details: {e}")
            return None

    def get_google_location_data(self, lat, lon):
        """Use Google's Maps API to collect location info for the provided latitude and longitude.

        Google returns a bunch of JSON with a variety of location data. This function returns
        Google's pre-formatted `formatted_address` key for a human-readable address.
        """
        google_maps_url = f"https://maps.googleapis.com/maps/api/geocode/json?latlng={lat},{lon}&sensor=false&key={self.GEOLOCATE_TOKEN}"
        r = requests.get(google_maps_url)
        maps_json = r.json()
        if r.ok:
            try:
                if "error_message" in maps_json:
                    print(f"[!] Google Maps returned an error so using Gophish coordinates. Error: {maps_json['error_message']}")
                    return f"{lat}, {lon}"
                first_result = maps_json['results'][0]
                if "formatted_address" in first_result:
                    return first_result["formatted_address"]
                # In case that key is ever unavailable try to assemble an address
                else:
                    components = first_result['address_components']
                    country = town = None
                    for c in components:
                        if "country" in c['types']:
                            country = c['long_name']
                        if "locality" in c['types']:
                            town = c['long_name']
                        if "administrative_area_level_1" in c['types']:
                            state = c['long_name']
                    return f"{town}, {state}, {country}"
            except Exception as e:
                print("[!] Failed to parse Google Maps API results so using Gophish coordinates.")
                print(f"L.. Error: {e}")
                return f"{lat}, {lon}"
        else:
            print(f"[!] Failed to contact the Google Maps API so using Gophish coordinates. Status code: {r.status_code}")
            return f"{lat}, {lon}"

    def geolocate(self, target, ipaddr, google=False):
        """Attempt to get location data for the provided target and event. Will use ipinfo.io if an
        API key is configured. Otherwise the Gophish latitude and longitude coordinates will be
        returned. If `google` is set to True this function will try to match the coordinates to a
        location using the Google Maps API.

        Returns a string: City, Region, Country
        """
        if ipaddr in self.ip_and_location:
            return self.ip_and_location[ipaddr]
        else:
            if self.IPINFO_TOKEN:
                # location_json = self.lookup_ip(event.details['browser']['address'])
                location_json = self.lookup_ip(ipaddr)
                if location_json:
                    city = region = country = "Unknown"
                    if "city" in location_json:
                        if location_json['city']:
                            city = location_json['city']
                    if "region" in location_json:
                        if location_json['region']:
                            region = location_json['region']
                    if "country" in location_json:
                        if location_json['country']:
                            country = location_json['country']
                    location = f"{city}, {region}, {country}"
                else:
                    location = f"{target.latitude}, {target.longitude}"
            elif google:
                if self.GEOLOCATE_TOKEN:
                    location = self.get_google_location_data(target.latitude, target.longitude)
                else:
                    location = f"{target.latitude}, {target.longitude}"
            else:
                location = f"{target.latitude}, {target.longitude}"
            self.locations.append(location)
            self.ip_and_location[ipaddr] = location
            return location

    def compare_ip_addresses(self, target_ip, browser_ip, verbose):
        """Compare the IP addresses of the target to that of an event. The goal: Looking for a
        mismatch that might identify some sort of interesting event. This might indicate an
        email was forwarded, a VPN was switched on/off, or maybe the target is at home.
        """
        if target_ip == browser_ip:
            return target_ip
        else:
            # We have an IP mismatch -- hard to tell why this might be.
            if verbose:
                print(f"[*] Event: This target's ({target_ip}) URL was clicked from a browser at {browser_ip}.")
            # This is an IP address not included in the results model, so we add it to our list here
            self.ip_addresses.append(browser_ip)
            return browser_ip

    def get_basic_campaign_info(self):
        """"Helper function to collect a campaign's basic details. This includes campaign name,
        status, template, and other details that are not the campaign's results.

        This keeps these calls in one place for tidiness and easier management.
        """
        self.cam_name = self.campaign.name
        self.cam_status = self.campaign.status
        self.created_date = self.campaign.created_date
        self.launch_date = self.campaign.launch_date
        self.completed_date = self.campaign.completed_date
        self.cam_url = self.campaign.url

        # Collect SMTP information
        self.smtp = self.campaign.smtp
        self.cam_from_address = self.smtp.from_address
        self.cam_smtp_host = self.smtp.host

        # Collect the template information
        self.template = self.campaign.template
        self.cam_subject_line = self.template.subject
        self.cam_template_name = self.template.name
        self.cam_template_attachments = self.template.attachments
        if self.cam_template_attachments == []:
            self.cam_template_attachments = "Non utilisé"

        # Collect the landing page information
        self.page = self.campaign.page
        self.cam_page_name = self.page.name
        self.cam_redirect_url = self.page.redirect_url
        if self.cam_redirect_url == "":
            self.cam_redirect_url = "Non utilisé"
        self.cam_capturing_passwords = self.page.capture_passwords
        self.cam_capturing_credentials = self.page.capture_credentials

    def collect_all_campaign_info(self, combine_reports):
        """Collect the campaign's details and set values for each of the variables."""
        # Collect the basic campaign details
        try:
            # Begin by checking if the ID is valid
            self.cam_id = self.campaign.id
            if combine_reports and self.cam_name is None:
                print(f"[+] Reports will be combined -- setting name, dates, and URL based on campaign ID {self.cam_id}.")
                self.get_basic_campaign_info()
            elif combine_reports is False:
                self.get_basic_campaign_info()
            # Collect the results and timeline lists
            if self.results is None:
                self.results = self.campaign.results
                self.timeline = self.campaign.timeline
            elif combine_reports:
                self.results += self.campaign.results
                self.timeline += self.campaign.timeline
            else:
                self.results = self.campaign.results
                self.timeline = self.campaign.timeline
        except:
            print(f"[!] Looks like campaign ID {self.cam_id} does not exist! Skipping it...")

    def process_results(self, combine_reports):
        """Process the results model to collect basic data, like total targets and event details.
        This should be run after the process_timeline_events() function which creates the
        targets_* lists.

        The results model can provide:
        first_name, last_name, email, position, and IP address
        """
        # Total length of results gives us the total number of targets
        if combine_reports and self.total_targets is None:
            self.total_targets = len(self.campaign.results)
        elif combine_reports:
            self.total_targets += len(self.campaign.results)
        else:
            # Not combining, so reset counters
            self.total_unique_opened = 0
            self.total_unique_clicked = 0
            self.total_unique_reported = 0
            self.total_unique_submitted = 0
            # Reports will not be combined, so reset tracking between reports
            self.total_targets = len(self.campaign.results)
            self.ip_addresses = []
            self.campaign_results_summary = []
        # Go through all results and extract data for statistics
        for target in self.campaign.results:
            temp_dict = {}
            # Log the IP address for additional statistics later
            if not target.ip == "":
                self.ip_addresses.append(target.ip)
                self.geolocate(target, target.ip, self.google)
            # Add all of the recipient's details and results to the temp dictionary
            temp_dict["email"] = target.email
            temp_dict["fname"] = target.first_name
            temp_dict["lname"] = target.last_name
            position = "None Provided"
            if target.position:
                position = target.position
            temp_dict["position"] = position
            temp_dict["ip_address"] = target.ip
            # Check if this target was recorded as viewing the email (tracking image)
            if target.email in self.targets_opened:
                temp_dict["opened"] = True
                self.total_unique_opened += 1
            else:
                temp_dict["opened"] = False
            # Check if this target clicked the link
            if target.email in self.targets_clicked:
                temp_dict["clicked"] = True
                self.total_unique_clicked += 1
                # Incremement the total number of opens for this target if they clicked
                # but did not display the tracking image in the email
                if target.email not in self.targets_opened:
                    self.total_unique_opened += 1
            else:
                temp_dict["clicked"] = False
            # Check if this target submitted data
            if target.email in self.targets_submitted:
                temp_dict["submitted"] = True
                self.total_unique_submitted += 1
            else:
                temp_dict["submitted"] = False
            # Check if this target reported the email
            if target.email in self.targets_reported:
                temp_dict["reported"] = True
                self.total_unique_reported += 1
            else:
                temp_dict["reported"] = False
            # Append the temp dictionary to the event summary list
            self.campaign_results_summary.append(temp_dict)

    def process_timeline_events(self, combine_reports):
        """Process the timeline model to collect basic data, like total clicks, and get detailed
        event data for recipients.

        The timeline model contains all events that occurred during the campaign.
        """
        # Create counters for enumeration
        sent_counter = 0
        click_counter = 0
        opened_counter = 0
        reported_counter = 0
        submitted_counter = 0

        # Reset target lists
        self.targets_opened = []
        self.targets_clicked = []
        self.targets_reported = []
        self.targets_submitted = []
        # Run through all events and count each of the four basic events
        for event in self.campaign.timeline:
            if event.message == "Email Sent":
                sent_counter += 1
            elif event.message == "Email Opened":
                opened_counter += 1
                self.targets_opened.append(event.email)
            elif event.message == "Clicked Link":
                click_counter += 1
                self.targets_clicked.append(event.email)
            elif event.message == "Submitted Data":
                submitted_counter += 1
                self.targets_submitted.append(event.email)
            elif event.message == "Email Reported":
                reported_counter += 1
                self.targets_reported.append(event.email)
        # Assign the counter values to our tracking lists
        if combine_reports:
            # Append, +=, totals if combining reports
            self.total_sent += sent_counter
            self.total_opened += opened_counter
            self.total_clicked += click_counter
            self.total_reported += reported_counter
            self.total_submitted += submitted_counter
        else:
            # Set tracking variables to current counter values for non-combined reports
            self.total_sent = sent_counter
            self.total_opened = opened_counter
            self.total_clicked = click_counter
            self.total_reported = reported_counter
            self.total_submitted = submitted_counter

    def generate_report(self):
        """Determines which type of report generate and the calls the appropriate reporting
        functions.
        """
        if self.report_format == "excel":
            print("[+] Building the report -- you selected a Excel/xlsx report.")
            self.output_xlsx_report = self._build_output_xlsx_file_name()
            self.write_xlsx_report()
        elif self.report_format == "word":
            template_path = os.path.join(os.path.dirname(__file__), "..", "templates", "template.docx")
            template_path = os.path.abspath(template_path)
            print("[+] Building the report -- you selected a Word/docx report.")
            print("[+] Looking for the template.docx to be used for the Word report.")
            if os.path.isfile(template_path):
                print("[+] Template was found -- proceeding with report generation...")
                print("L.. Word reports can take a while if you had a lot of recipients.")
                self.output_word_report = self._build_output_word_file_name()
                self.write_word_report(template_path)
            else:
                print("[!] Could not find the template document! Make sure 'template.docx' is in the GoReport directory.")
                sys.exit()
        elif self.report_format == "quick":
            print("[+] Quick report stats:")
            self.get_quick_stats()

    def get_quick_stats(self):
        """Present quick stats for the campaign. Just basic numbers and some details."""
        print()
        print(self.cam_name)
        print(f"Status:\t\t{self.cam_status}")
        print(f"Created:\t{self.created_date.split('T')[1].split('.')[0]} on {self.created_date.split('T')[0]}")
        print(f"Started:\t{self.launch_date.split('T')[1].split('.')[0]} on {self.launch_date.split('T')[0]}")
        if self.cam_status == "Completed":
            print(f"Completed:\t{self.completed_date.split('T')[1].split('.')[0]} on {self.completed_date.split('T')[0]}")
        print()
        print(f"Total Targets:\t{self.total_targets}")
        print(f"Emails Sent:\t{self.total_sent}")
        print(f"IPs Seen:\t{len(self.ip_addresses)}")
        print()
        print(f"Total Opened Events:\t\t{self.total_opened}")
        print(f"Total Click Events:\t\t{self.total_clicked}")
        print(f"Total Submitted Data Events:\t{self.total_submitted}")
        print()
        print(f"Individuals Who Opened:\t\t\t{self.total_unique_opened}")
        print(f"Individuals Who Clicked:\t\t{self.total_unique_clicked}")
        print(f"Individuals Who Entered Data:\t\t{self.total_unique_submitted}")
        print(f"Individuals Who Reported the Email:\t{self.total_unique_reported}")

    def _build_output_xlsx_file_name(self):
        """Create the xlsx report name and save it to ../reports/"""
        safe_name = "".join([c for c in self.cam_name if c.isalpha() or c.isdigit() or c == " "]).rstrip()
        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        filename = f"rapport_campagne_{self.cam_id}_{timestamp}.xlsx"
        reports_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "reports"))
        os.makedirs(reports_dir, exist_ok=True)
        return os.path.join(reports_dir, filename)

    def _build_output_word_file_name(self):
        """Create the docx report name and save it to ../reports/"""
        safe_name = "".join([c for c in self.cam_name if c.isalpha() or c.isdigit() or c == " "]).rstrip()
        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        filename = f"rapport_campagne_{self.cam_id}_{timestamp}.docx"
        reports_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "reports"))
        os.makedirs(reports_dir, exist_ok=True)
        return os.path.join(reports_dir, filename)

    def _set_word_column_width(self, column, width):
        """Custom function for quickly and easily setting the width of a table's column in the Word
        docx output.

        This option is missing from the basic Python-docx library.
        """
        for cell in column.cells:
            cell.width = width

    def write_xlsx_report(self):
        """Assemble and output the xlsx file report.

        Throughout this function, results are assembled by adding commas and then adding to a
        results string, i.e. 'result_A' and then 'result_A' += ',result_B'. This is so the
        result can be written to the csv file and have the different pieces end up in the correct
        columns.
        """
        goreport_xlsx = xlsxwriter.Workbook(self.output_xlsx_report)
        # Bold format
        bold_format = goreport_xlsx.add_format({'bold': True})
        bold_format.set_text_wrap()
        bold_format.set_align('vcenter')
        # Centered format
        center_format = goreport_xlsx.add_format()
        center_format.set_text_wrap()
        center_format.set_align('vcenter')
        center_format.set_align('center')
        # Header format
        header_format = goreport_xlsx.add_format({'bold': True})
        header_format.set_text_wrap()
        header_format.set_align('vcenter')
        header_format.set_bg_color(self.xlsx_header_bg_color)
        header_format.set_font_color(self.xlsx_header_font_color)
        # Number cells
        num_format = goreport_xlsx.add_format()
        num_format.set_align('center')
        # Boolean cells - True
        true_format = goreport_xlsx.add_format({'bold': True})
        true_format.set_text_wrap()
        true_format.set_align('vcenter')
        true_format.set_font_color("#9C0006")
        true_format.set_bg_color("#FFC7CE")
        # Boolean cells - True
        false_format = goreport_xlsx.add_format()
        false_format.set_text_wrap()
        false_format.set_align('vcenter')
        false_format.set_font_color("#006100")
        false_format.set_bg_color("#C6EFCE")
        # Remaining cells
        wrap_format = goreport_xlsx.add_format()
        wrap_format.set_text_wrap()
        wrap_format.set_align('vcenter')

        worksheet = goreport_xlsx.add_worksheet("Overview")
        col = 0
        row = 0

        worksheet.set_column(0, 10, 62)

        worksheet.write(row, col, "Campaign Results For:", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_name}", wrap_format)
        row += 1
        worksheet.write(row, col, "Status", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_status}", wrap_format)
        row += 1
        worksheet.write(row, col, "Created", bold_format)
        worksheet.write(row, col + 1, f"{self.created_date}", wrap_format)
        row += 1
        worksheet.write(row, col, "Started", bold_format)
        worksheet.write(row, col + 1, f"{self.launch_date}", wrap_format)
        row += 1
        if self.cam_status == "Completed":
            worksheet.write(row, col, "Completed", bold_format)
            worksheet.write(row, col + 1, f"{self.completed_date}", wrap_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Campaign Details", bold_format)
        row += 1
        worksheet.write(row, col, "From", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_from_address}", wrap_format)
        row += 1
        worksheet.write(row, col, "Subject", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_subject_line}", wrap_format)
        row += 1
        worksheet.write(row, col, "Phish URL", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_url}", wrap_format)
        row += 1
        worksheet.write(row, col, "Redirect URL", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_redirect_url}", wrap_format)
        row += 1
        worksheet.write(row, col, "Attachment(s)", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_template_attachments}", wrap_format)
        row += 1
        worksheet.write(row, col, "Captured Passwords", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_capturing_credentials}", wrap_format)
        row += 1
        worksheet.write(row, col, "Stored Passwords", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_capturing_passwords}", wrap_format)
        row += 1

        worksheet.write(row, col, "")
        row += 1

        # Write a high level summary for stats
        worksheet.write(row, col, "High Level Results", bold_format)
        row += 1
        worksheet.write(row, col, "Total Targets", bold_format)
        worksheet.write(row, col + 1, self.total_targets, num_format)
        row += 1

        worksheet.write(row, col, "The following totals indicate how many events of each type Gophish recorded:", wrap_format)
        row += 1
        worksheet.write(row, col, "Total Opened Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_opened, num_format)
        row += 1
        worksheet.write(row, col, "Total Clicked Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_clicked, num_format)
        row += 1
        worksheet.write(row, col, "Total Submitted Data Events", bold_format)
        worksheet.write(row, col + 1, "", wrap_format)
        row += 1
        worksheet.write(row, col, "Total Report Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_reported, num_format)
        row += 1

        worksheet.write(row, col, "The following totals indicate how many targets participated in each event type:", wrap_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Opened", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_opened, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Clicked", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_clicked, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Submitted Data", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_submitted, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Reported", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_reported, num_format)
        row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet = goreport_xlsx.add_worksheet("Summary")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 20)

        worksheet.write(row, col, "Summary of Events", bold_format)
        row += 1

        header_col = 0
        headers = ["Email Address", "Open", "Click", "Creds", "Report", "OS", "Browser"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1

        # Sort campaign summary by each dict's email entry and then create results table
        target_counter = 0
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
        for target in ordered_results:
            worksheet.write(row, col, target['email'], wrap_format)
            if target['opened']:
                worksheet.write_boolean(row, col + 1, target['opened'], true_format)
            else:
                worksheet.write_boolean(row, col + 1, target['opened'], false_format)
            if target['clicked']:
                worksheet.write_boolean(row, col + 2, target['clicked'], true_format)
            else:
                worksheet.write_boolean(row, col + 2, target['clicked'], false_format)
            if target['submitted']:
                worksheet.write_boolean(row, col + 3, target['submitted'], true_format)
            else:
                worksheet.write_boolean(row, col + 3, target['submitted'], false_format)
            if target['reported']:
                worksheet.write_boolean(row, col + 4, target['reported'], true_format)
            else:
                worksheet.write_boolean(row, col + 4, target['reported'], false_format)
            if target['email'] in self.targets_clicked:
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 5, browser_details, wrap_format)
                        worksheet.write(row, col + 6, os_details, wrap_format)
            else:
                worksheet.write(row, col + 5, "N/A", wrap_format)
                worksheet.write(row, col + 6, "N/A", wrap_format)
            row += 1
            target_counter += 1
            print(f"[+] Created row for {target_counter} of {self.total_targets}.")

        print("[+] Finished writing events summary...")
        print("[+] Detailed results analysis is next and will take some time if you had a lot of targets...")
        # End of the event summary and beginning of the detailed results

        worksheet = goreport_xlsx.add_worksheet("Event Details")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 40)

        worksheet.write(row, col, "Detailed Analysis", bold_format)
        row += 1

        target_counter = 0
        for target in self.results:
            # Only create a Detailed Analysis section for targets with clicks
            if target.email in self.targets_clicked:
                position = ""
                if target.position:
                    position = f"({target.position})"
                worksheet.write(row, col, f"{target.first_name} {target.last_name} {position}", bold_format)
                row += 1
                worksheet.write(row, col, target.email, wrap_format)
                row += 1
                # Go through all events to find events for this target
                for event in self.timeline:
                    if event.message == "Email Sent" and event.email == target.email:
                        # Parse the timestamp into separate date and time variables
                        temp = event.time.split('T')
                        sent_date = temp[0]
                        sent_time = temp[1].split('.')[0]
                        # Record the email sent date and time in the report
                        worksheet.write(row, col, f"Sent on {sent_date.replace(',', '')} at {sent_time}", wrap_format)
                        row += 1

                    if event.message == "Email Opened" and event.email == target.email:
                        # Record the email preview date and time in the report
                        temp = event.time.split('T')
                        worksheet.write(row, col, f"Email Preview at {temp[0]} {temp[1].split('.')[0]}", wrap_format)
                        row += 1

                    if event.message == "Clicked Link" and event.email == target.email:
                        worksheet.write(row, col, "Email Link Clicked", bold_format)
                        row += 1

                        header_col = 0
                        headers = ["Time", "IP", "Location", "Browser", "Operating System"]
                        for header in headers:
                            worksheet.write(row, header_col, header, header_format)
                            header_col += 1
                        row += 1

                        temp = event.time.split('T')
                        worksheet.write(row, col, f"{temp[0]} {temp[1].split('.')[0]}", wrap_format)

                        # Check if browser IP matches the target's IP and record result
                        ip_comparison = self.compare_ip_addresses(target.ip,
                                                                  event.details['browser']['address'],
                                                                  self.verbose)
                        worksheet.write(row, col + 1, f"{ip_comparison}", wrap_format)

                        # Parse the location data
                        loc = self.geolocate(target, event.details['browser']['address'], self.google)
                        worksheet.write(row, col + 2, loc, wrap_format)

                        # Parse the user-agent string and add browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        worksheet.write(row, col + 3, browser_details, wrap_format)
                        self.browsers.append(browser_details)

                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 4, os_details, wrap_format)
                        self.operating_systems.append(os_details)
                        row += 1

                    if event.message == "Submitted Data" and event.email == target.email:
                        # Now we have events for submitted data. A few notes on this:
                        #   1. There is no expectation of a Submit event without a Clicked Link event
                        #   2. Assuming that, the following process does NOT flag IP mismatches
                        #      or add to the list of seen locations, OSs, IPs, or browsers.
                        worksheet.write(row, col, "Submitted Data Captured", bold_format)
                        row += 1

                        header_col = 0
                        headers = ["Time", "IP", "Location", "Browser", "Operating System", "Data Captured"]
                        for header in headers:
                            worksheet.write(row, header_col, header, header_format)
                            header_col += 1
                        row += 1

                        temp = event.time.split('T')
                        worksheet.write(row, col, f"{temp[0]} {temp[1].split('.')[0]}", wrap_format)

                        worksheet.write(row, col + 1, f"{event.details['browser']['address']}", wrap_format)

                        loc = self.geolocate(target, event.details['browser']['address'], self.google)
                        worksheet.write(row, col + 2, loc, wrap_format)

                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        worksheet.write(row, col + 3, browser_details, wrap_format)

                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 4, os_details, wrap_format)

                        # Get just the submitted data from the event's payload
                        submitted_data = ""
                        data_payload = event.details['payload']
                        # Get all of the submitted data
                        for key, value in data_payload.items():
                            # To get just submitted data, we drop the 'rid' key
                            if not key == "rid":
                                submitted_data += f"{key}:{str(value).strip('[').strip(']')}"
                        worksheet.write(row, col + 5, submitted_data, wrap_format)
                        row += 1

                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
            else:
                # This target had no clicked or submitted events so move on to next
                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
                continue
            worksheet.write(row, col, "")
            row += 1

        print("[+] Finished writing detailed analysis...")

        worksheet = goreport_xlsx.add_worksheet("Stats")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 35)

        worksheet.write(row, col, "Recorded Browsers Based on User-Agents:", bold_format)
        row += 1

        header_col = 0
        headers = ["Browser", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_browsers = Counter(self.browsers)
        for key, value in counted_browsers.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Record OS From Browser User-Agents:", bold_format)
        row += 1
        header_col = 0
        headers = ["Operating System", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_os = Counter(self.operating_systems)
        for key, value in counted_os.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Recorded Locations from IPs:", bold_format)
        row += 1
        header_col = 0
        headers = ["Locations", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_locations = Counter(self.locations)
        for key, value in counted_locations.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Recorded IPs:", bold_format)
        row += 1
        header_col = 0
        headers = ["IP Address", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_ip_addresses = Counter(self.ip_addresses)
        for key, value in counted_ip_addresses.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "Recorded IPs and Locations:", bold_format)
        row += 1
        header_col = 0
        headers = ["IP Address", "Location"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        for key, value in self.ip_and_location.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write(row, col + 1, f"{value}", wrap_format)
            row += 1

        goreport_xlsx.close()
        print(f"[+] Done! Check '{self.output_xlsx_report}' for your results.")

    def write_word_report(self, template_path):
        """Assemble and output the Word docx file report."""
        # Create document writer using the template and a style editor
        d = Document(template_path)
        styles = d.styles
        print(template_path)

        # Create a custom styles for table cells
        _ = styles.add_style("Cell Text", WD_STYLE_TYPE.CHARACTER)
        cell_text = d.styles["Cell Text"]
        cell_text_font = cell_text.font
        cell_text_font.name = "Calibri"
        cell_text_font.size = Pt(12)
        cell_text_font.bold = True
        cell_text_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        _ = styles.add_style("Cell Text Hit", WD_STYLE_TYPE.CHARACTER)
        cell_text_hit = d.styles["Cell Text Hit"]
        cell_text_hit_font = cell_text_hit.font
        cell_text_hit_font.name = "Calibri"
        cell_text_hit_font.size = Pt(12)
        cell_text_hit_font.bold = True
        cell_text_hit_font.color.rgb = RGBColor(0x00, 0x96, 0x00)

        _ = styles.add_style("Cell Text Miss", WD_STYLE_TYPE.CHARACTER)
        cell_text_miss = d.styles["Cell Text Miss"]
        cell_text_miss_font = cell_text_miss.font
        cell_text_miss_font.name = "Calibri"
        cell_text_miss_font.size = Pt(12)
        cell_text_miss_font.bold = True
        cell_text_miss_font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Write a campaign summary at the top of the report
        d.add_heading("Résumé Exécutif", 1)
        p = d.add_paragraph()
        run = p.add_run(f"Résultats de la campagne pour: {self.cam_name}")
        run.bold = True
        # Runs are basically "runs" of text and must be aligned like we want
        # them aligned in the report -- thus they are pushed left
        if self.cam_status == "Completed":
            completed_status = f"{self.completed_date.split('T')[1].split('.')[0]} le {self.completed_date.split('T')[0]}"
        else:
            completed_status = "En cours"
        p.add_run(f"""
Statut: {self.cam_status}
Créé à: {self.created_date.split('T')[1].split('.')[0]} le {self.created_date.split('T')[0]}
Début à: {self.launch_date.split('T')[1].split('.')[0]} le {self.launch_date.split('T')[0]}
Terminé à: {completed_status}

""")
        if self.cam_status == "Completed":
            print()

        # Write the campaign details -- email details and template settings
        run = p.add_run("Détails de la campagne")
        run.bold = True
        p.add_run(f"""
De: {self.cam_from_address}
Objet: {self.cam_subject_line}
URL de phishing: {self.cam_url}
URL de redirection: {self.cam_redirect_url}
Pièce(s) jointe(s): {self.cam_template_attachments}
Identifiants capturés: {"Oui" if self.cam_capturing_credentials else "Non"}
Mots de passe stockés: {"Oui" if self.cam_capturing_passwords else "Non"}

""")

        # Write a high level summary for stats
        run = p.add_run("Résultats Globaux")
        run.bold = True
        p.add_run(f"""
Nombre total de destinataires: {self.total_targets}

Les résultats suivants indiquent le nombre d'événements de chaque type enregistrés par Gophish:
Total d'événements d'ouverture: {self.total_opened}
Total d'événements de clic: {self.total_clicked}
Total d'événements de signalement: {self.total_reported}
Total d'événements de soumission de données: {self.total_submitted}

Les totaux suivants indiquent le nombre de destinataires ayant participé à chaque type d'événement:
Personnes ayant ouvert: {self.total_unique_opened}
Personnes ayant cliqué: {self.total_unique_clicked}
Personnes ayant signalé l'email: {self.total_unique_reported}
Personnes ayant soumis des données: {self.total_unique_submitted}

""")
        d.add_page_break()

        print("[+] Finished writing high level summary...")
        # End of the campaign summary and beginning of the event summary
        d.add_heading("Résumé des Événements", 1)
        d.add_paragraph("Le tableau suivant résume qui a ouvert et cliqué sur les emails envoyés dans cette campagne.")

        # Create a table to hold the event summary results
        table = d.add_table(rows=len(self.campaign_results_summary) + 1, cols=7, style="GoReport")

        header0 = table.cell(0, 0)
        header0.text = ""
        header0.paragraphs[0].add_run("Adresse email", "Cell Text").bold = True

        header1 = table.cell(0, 1)
        header1.text = ""
        header1.paragraphs[0].add_run("Ouvert", "Cell Text").bold = True

        header2 = table.cell(0, 2)
        header2.text = ""
        header2.paragraphs[0].add_run("Clic", "Cell Text").bold = True

        header3 = table.cell(0, 3)
        header3.text = ""
        header3.paragraphs[0].add_run("Données", "Cell Text").bold = True

        header4 = table.cell(0, 4)
        header4.text = ""
        header4.paragraphs[0].add_run("Signalement", "Cell Text").bold = True

        header5 = table.cell(0, 5)
        header5.text = ""
        header5.paragraphs[0].add_run("OS", "Cell Text").bold = True

        header6 = table.cell(0, 6)
        header6.text = ""
        header6.paragraphs[0].add_run("Navigateur", "Cell Text").bold = True

        # Sort campaign summary by each dict's email entry and then create results table
        target_counter = 0
        counter = 1
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
        for target in ordered_results:
            email_cell = table.cell(counter, 0)
            email_cell.text = f"{target['email']}"

            temp_cell = table.cell(counter, 1)
            if target['opened']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 2)
            if target['clicked']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 3)
            if target['submitted']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 4)
            if target['reported']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            if target['email'] in self.targets_clicked:
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        os_details = user_agent.os.family + " " + \
                            user_agent.os.version_string
                        temp_cell = table.cell(counter, 5)
                        temp_cell.text = os_details
                        temp_cell = table.cell(counter, 6)
                        temp_cell.text = browser_details
            else:
                temp_cell = table.cell(counter, 5)
                temp_cell.text = "N/A"
                temp_cell = table.cell(counter, 6)
                temp_cell.text = "N/A"
            counter += 1
            target_counter += 1
            print(f"[+] Created table entry for {target_counter} of {self.total_targets}.")

        d.add_page_break()

        # End of the event summary and beginning of the detailed results
        print("[+] Finished writing events summary...")
        print("[+] Detailed results analysis is next and may take some time if you had a lot of targets...")
        d.add_heading("Analyse Détaillée", 1)
        sofware_infos_dict = {} ### ADDED
        WRITE_CVE_PARTS = True ### ADDED
        target_counter = 0
        for target in self.results:
            # Only create a Detailed Analysis section for targets with clicks
            if target.email in self.targets_clicked:
                # Create counters to track table cell locations
                opened_counter = 1
                clicked_counter = 1
                submitted_counter = 1
                # Create section starting with a header with the first and last name
                position = ""
                if target.position:
                    position = f"({target.position})"
                d.add_heading(f"{target.first_name} {target.last_name} {position}", 2)
                p = d.add_paragraph(target.email)
                # p = d.add_paragraph()
                # Save a spot to record the email sent date and time in the report
                email_sent_run = p.add_run()
                # Go through all events to find events for this target
                ### ADDED
                clicked_user_agent = None
                ### END ADDED
                previous_category = False
                for event in self.timeline:
                    if event.message == "Email Sent" and event.email == target.email:
                        # Parse the timestamp into separate date and time variables
                        # Ex: 2017-01-30T14:31:22.534880731-05:00
                        temp = event.time.split('T')
                        sent_date = temp[0]
                        sent_time = temp[1].split('.')[0]
                        # Record the email sent date and time in the run created earlier
                        email_sent_run.text = f" - Email envoyé le {sent_date} à {sent_time}"
                        previous_category = True
                    if event.message == "Email Opened" and event.email == target.email:
                        if opened_counter == 1:
                            # Create the Email Opened/Previewed table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            if previous_category:
                                _ = p.add_run()
                                _.text = "\n"
                            else:
                                previous_category = True
                            run = p.add_run("Aperçu des courriels")
                            run.bold = True

                            opened_table = d.add_table(rows=1, cols=1, style="GoReport")
                            opened_table.autofit = True
                            opened_table.allow_autofit = True

                            header1 = opened_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Heure", "Cell Text").bold = True

                        # Begin by adding a row to the table and inserting timestamp
                        opened_table.add_row()
                        timestamp = opened_table.cell(opened_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]
                        opened_counter += 1

                    if event.message == "Clicked Link" and event.email == target.email:
                        if clicked_counter == 1:
                            # Create the Clicked Link table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            if previous_category:
                                _ = p.add_run()
                                _.text = "\n"
                            else:
                                previous_category = True
                            run = p.add_run("Lien de l'email cliqué")
                            run.bold = True

                            clicked_table = d.add_table(rows=1, cols=5, style="GoReport")
                            clicked_table.autofit = True
                            clicked_table.allow_autofit = True

                            header1 = clicked_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Heure", "Cell Text").bold = True

                            header2 = clicked_table.cell(0, 1)
                            header2.text = ""
                            header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

                            header3 = clicked_table.cell(0, 2)
                            header3.text = ""
                            header3.paragraphs[0].add_run("Localisation", "Cell Text").bold = True

                            header4 = clicked_table.cell(0, 3)
                            header4.text = ""
                            header4.paragraphs[0].add_run("Navigateur", "Cell Text").bold = True

                            header5 = clicked_table.cell(0, 4)
                            header5.text = ""
                            header5.paragraphs[0].add_run("Système d'exploitation (OS)",
                                                          "Cell Text").bold = True

                        clicked_table.add_row()
                        timestamp = clicked_table.cell(clicked_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        ip_add = clicked_table.cell(clicked_counter, 1)
                        # Check if browser IP matches the target's IP and record result
                        ip_add.text = self.compare_ip_addresses(
                            target.ip, event.details['browser']['address'], self.verbose)

                        # Parse the location data
                        event_location = clicked_table.cell(clicked_counter, 2)
                        event_location.text = self.geolocate(target, event.details['browser']['address'], self.google)

                        # Parse the user-agent string for browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        browser = clicked_table.cell(clicked_counter, 3)
                        browser.text = browser_details
                        self.browsers.append(browser_details)
                        
                        clicked_user_agent = event.details['browser']['user-agent']

                        op_sys = clicked_table.cell(clicked_counter, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = os_details
                        self.operating_systems.append(os_details)

                        clicked_counter += 1

                    if event.message == "Submitted Data" and event.email == target.email:
                        if submitted_counter == 1:
                            # Create the Submitted Data table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            if previous_category:
                                _ = p.add_run()
                                _.text = "\n"
                            else:
                                previous_category = True
                            run = p.add_run("Données capturées")
                            run.bold = True

                            submitted_table = d.add_table(rows=1, cols=6, style="GoReport")
                            submitted_table.autofit = True
                            submitted_table.allow_autofit = True

                            header1 = submitted_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Heure", "Cell Text").bold = True

                            header2 = submitted_table.cell(0, 1)
                            header2.text = ""
                            header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

                            header3 = submitted_table.cell(0, 2)
                            header3.text = ""
                            header3.paragraphs[0].add_run("Localisation", "Cell Text").bold = True

                            header4 = submitted_table.cell(0, 3)
                            header4.text = ""
                            header4.paragraphs[0].add_run("Navigateur", "Cell Text").bold = True

                            header5 = submitted_table.cell(0, 4)
                            header5.text = ""
                            header5.paragraphs[0].add_run("Système d'exploitation (OS)",
                                                          "Cell Text").bold = True

                            header6 = submitted_table.cell(0, 5)
                            header6.text = ""
                            header6.paragraphs[0].add_run("Données capturées",
                                                          "Cell Text").bold = True

                        submitted_table.add_row()
                        timestamp = submitted_table.cell(submitted_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        ip_add = submitted_table.cell(submitted_counter, 1)
                        ip_add.text = event.details['browser']['address']

                        # Parse the location data
                        event_location = submitted_table.cell(submitted_counter, 2)
                        event_location.text = self.geolocate(target, event.details['browser']['address'], self.google)

                        # Parse the user-agent string and add browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        browser = submitted_table.cell(submitted_counter, 3)
                        browser.text = browser_details

                        op_sys = submitted_table.cell(submitted_counter, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = f"{os_details}"

                        # Get just the submitted data from the event's payload
                        submitted_data = ""
                        data = submitted_table.cell(submitted_counter, 5)
                        data_payload = event.details['payload']
                        # Get all of the submitted data
                        for key, value in data_payload.items():
                            # To get just submitted data, we drop the 'rid' key
                            if not key == "rid":
                                submitted_data += f"{key}:{str(value).strip('[').strip(']')}   "
                        data.text = f"{submitted_data}"
                        submitted_counter += 1
                        
                #########################
                ##### DEBUT CUSTOM ######
                #########################
                
                # Create the Common Vulnerabilities and Exposures
                if clicked_user_agent and WRITE_CVE_PARTS:
                    p = d.add_paragraph()
                    p.style = d.styles['Normal']
                    _ = p.add_run()
                    _.text = "\n"
                    run = p.add_run("Vulnérabilités et Expositions Courantes (CVE)")
                    run.bold = True

                    try:
                        software_table = d.add_table(rows=2, cols=3, style="GoReport")
                        software_table.autofit = True
                        software_table.allow_autofit = True

                        header1 = software_table.cell(0, 0)
                        header1.text = ""
                        header1.paragraphs[0].add_run("Navigateur / OS utilisé", "Cell Text").bold = True

                        header2 = software_table.cell(0, 1)
                        header2.text = ""
                        header2.paragraphs[0].add_run("Dernières versions", "Cell Text").bold = True

                        header3 = software_table.cell(0, 2)
                        header3.text = ""
                        header3.paragraphs[0].add_run("À jour", "Cell Text").bold = True

                        cve_infos = {
                            'browser_cpe': None,
                            'os_cpe': None,
                            'last_version_browser': None,
                            'last_version_os': None,
                            'version_browser_outdated': None,
                            'version_os_outdated': None,
                            'vulnerable_date_browser': None,
                            'vulnerable_date_os': None,
                            'vulnerabilities': None,
                            'high_severity_count': None,
                            'critical_severity_count': None,
                            'most_impactful_vuln': None,
                        }
                        try:
                            cve_infos = custom_cve.search_user_agent_vulnerable(clicked_user_agent)
                            # fake_ua = "Mozilla/5.0 (iPhone; CPU iPhone OS 17_7_2 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) FxiOS/119.0 Mobile/15E148 Safari/605.1.15"
                            # cve_infos = custom_cve.search_user_agent_vulnerable([fake_ua, clicked_user_agent][target_counter % 2])
                        except Exception as e:
                            pass
                        
                        software_table.add_row()
                        # software_table.format.fit = 'auto'
                        used_browser = software_table.cell(1, 0)
                        # Parse the user-agent string for browser and OS details
                        parsed_user_agent = parse(clicked_user_agent)
                        browser_details = parsed_user_agent.browser.family + " " + \
                            parsed_user_agent.browser.version_string
                        used_browser.text = browser_details

                        latest_browser = software_table.cell(1, 1)
                        latest_browser_formatted = parsed_user_agent.browser.family + " " + \
                            cve_infos.get('last_version_browser', 'N/A')
                        latest_browser.text = latest_browser_formatted
                            
                        sofware_infos_dict[browser_details] = cve_infos
                        sofware_infos_dict[browser_details]['latest_formatted'] = latest_browser_formatted
                        
                        uptodate = software_table.cell(1, 2)
                        # Add Hit or Miss to the table cell
                        if cve_infos.get('version_browser_outdated', False):
                            uptodate.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")
                        else:
                            uptodate.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
                        
                        used_os = software_table.cell(2, 0)
                        os_details = parsed_user_agent.os.family + " " + parsed_user_agent.os.version_string
                        used_os.text = os_details
                        
                        latest_os = software_table.cell(2, 1)
                        latest_os_formatted = parsed_user_agent.os.family + " " + \
                            cve_infos.get('last_version_os', 'N/A')
                        latest_os.text = latest_os_formatted
                        
                        sofware_infos_dict[os_details] = cve_infos
                        sofware_infos_dict[os_details]['latest_formatted'] = latest_os_formatted
                            
                        uptodate = software_table.cell(2, 2)
                        # Add Hit or Miss to the table cell
                        if cve_infos.get('version_os_outdated', False):
                            uptodate.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")
                        else:
                            uptodate.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
                        
                        p = d.add_paragraph()
                        # p.style = d.styles['Normal']
                        # run = p.add_run("\n")
                        
                        one_vulnerable = False
                        
                        # If browser is vulnerable, add a text 'Browser is vulnerable since ..'
                        vulnerable_date = cve_infos.get('vulnerable_date_browser', False)
                        # from datetime import timedelta
                        # vulnerable_date = datetime.now() - timedelta(days=43)
                        if vulnerable_date:
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run(f"Le navigateur est vulnérable depuis ")
                            run = p.add_run(f"{vulnerable_date.strftime('%Y-%m-%d')} ({(datetime.now() - vulnerable_date).days} jours)")
                            run.bold = True
                            one_vulnerable = True
                        
                        # If OS is vulnerable, add a text 'OS is vulnerable since ..'
                        vulnerable_date = cve_infos.get('vulnerable_date_os', False)
                        # vulnerable_date = datetime.now() - timedelta(days=345)
                        if vulnerable_date:
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run(f"Le système d'exploitation est vulnérable depuis ")
                            run = p.add_run(f"{vulnerable_date.strftime('%Y-%m-%d')} ({(datetime.now() - vulnerable_date).days} jours)")
                            run.bold = True
                            one_vulnerable = True
                        
                        if one_vulnerable:
                            run = p.add_run("\n")
                        
                        # Show a list of the X first vulnerabilities in a table
                        # p = d.add_paragraph()
                        # p.style = d.styles['Normal']
                        # run = p.add_run("Vulnerabilities")
                        # run.bold = True
                        
                        vulnerabilities = cve_infos.get('vulnerabilities', [])
                        
                        if len(vulnerabilities) > 0:
                            vulnerabilities_table = d.add_table(rows=1, cols=4, style="GoReport")
                            vulnerabilities_table.autofit = True
                            vulnerabilities_table.allow_autofit = True
                            
                            header1 = vulnerabilities_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("CVE ID", "Cell Text").bold = True
                            
                            header2 = vulnerabilities_table.cell(0, 1)
                            header2.text = ""
                            header2.paragraphs[0].add_run("Date", "Cell Text").bold = True
                            
                            header3 = vulnerabilities_table.cell(0, 2)
                            header3.text = ""
                            header3.paragraphs[0].add_run("CVSS Score", "Cell Text").bold = True
                            
                            header4 = vulnerabilities_table.cell(0, 3)
                            header4.text = ""
                            header4.paragraphs[0].add_run("Gravité", "Cell Text").bold = True
                            
                            # template : vulnerabiliy = 
                            # {
                            #     'id': cve_id,
                            #     'published_date': published_date,
                            #     'description': description,
                            #     'highest_score': highest_score,
                            #     'highest_severity': highest_severity,
                            #     'specific_match': specific_match,
                            #     'impacts': impacts,
                            # }
                            
                            # high_severity_bg_color = RGBColor(0xF7, 0x80, 0x70) # red
                            high_severity_font_color = RGBColor(0xFF, 0x00, 0x00) # red
                            # critical_severity_bg_color = RGBColor(0x00, 0x00, 0x00) # black
                            critical_severity_font_color = RGBColor(0x00, 0x00, 0x00) # black
                            
                            cvss_normal = RGBColor(0xFF, 0x00, 0x00)
                            cvss_more_than_9 = RGBColor(0x00, 0x00, 0x00)
                            
                            def add_hyperlink(paragraph, url, text, color, underline):
                                # This gets access to the document.xml.rels file and gets a new relation id value
                                part = paragraph.part
                                r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                                # Create the w:hyperlink tag and add needed values
                                hyperlink = oxml.shared.OxmlElement('w:hyperlink')
                                hyperlink.set(oxml.shared.qn('r:id'), r_id, )
                                # Create a w:r element
                                new_run = oxml.shared.OxmlElement('w:r')
                                # Create a new w:rPr element
                                rPr = oxml.shared.OxmlElement('w:rPr')
                                # Add color if it is given
                                if not color is None:
                                    c = oxml.shared.OxmlElement('w:color')
                                    c.set(oxml.shared.qn('w:val'), color)
                                    rPr.append(c)
                                # Remove underlining if it is requested
                                if underline:
                                    u = oxml.shared.OxmlElement('w:u')
                                    u.set(oxml.shared.qn('w:val'), 'single')
                                    rPr.append(u)
                                else:
                                    u = oxml.shared.OxmlElement('w:u')
                                    u.set(oxml.shared.qn('w:val'), 'none')
                                    rPr.append(u)
                                # Join all the xml elements together  add the required text to the w:r element
                                new_run.append(rPr)
                                new_run.text = text
                                hyperlink.append(new_run)
                                paragraph._p.append(hyperlink)
                                return hyperlink
                            
                            CVE_URL = "https://nvd.nist.gov/vuln/detail/"
                            
                            i = 0
                            max_items = 5
                            for cve in vulnerabilities:
                                if i < max_items:
                                    vulnerabilities_table.add_row()
                                    cve_id = vulnerabilities_table.cell(i + 1, 0)
                                    cve_id_str = cve.get('id', 'N/A')
                                    # cve_id.text = cve_id_str
                                    # Add the CVE ID to the cell with a hyperlink
                                    cve_id_paragraph = cve_id.paragraphs[0]
                                    hyperlink = add_hyperlink(cve_id_paragraph, CVE_URL + cve_id_str, cve_id_str, '0000FF', True)
                                    
                                    date = vulnerabilities_table.cell(i + 1, 1)
                                    published_date = cve.get('published_date', None)
                                    if published_date:
                                        published_date = published_date.strftime('%Y-%m-%d')
                                    else:
                                        published_date = "N/A"
                                    date.text = published_date
                                    
                                    cvss_score = vulnerabilities_table.cell(i + 1, 2)
                                    cvss_score_float = float(cve.get('highest_score', -1.0))
                                    cvss_score_str = str(cvss_score_float)
                                    run = cvss_score.paragraphs[0].add_run(cvss_score_str)
                                    if cvss_score_float > 9.0:
                                        run.font.color.rgb = cvss_more_than_9
                                        run.bold = True
                                    else:
                                        run.font.color.rgb = cvss_normal
                                        run.bold = True
                                    
                                    severity = vulnerabilities_table.cell(i + 1, 3)
                                    severity_str = cve.get('highest_severity', 'N/A')
                                    run = severity.paragraphs[0].add_run(severity_str)
                                    if severity_str.lower() == "high":
                                        run.font.color.rgb = high_severity_font_color
                                        run.bold = True
                                        # severity.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.RED  # Red background
                                    elif severity_str.lower() == "critical":
                                        run.font.color.rgb = critical_severity_font_color
                                        run.bold = True
                                else:
                                    # Add a '+ XX more' line
                                    vulnerabilities_table.add_row()
                                    cve_id = vulnerabilities_table.cell(i + 1, 0)
                                    run = cve_id.paragraphs[0].add_run(f"+ {len(cve_infos.get('vulnerabilities', [])) - max_items} autres")
                                    run.bold = True
                                    run.italic = True
                                    break
                                i += 1
                            
                            # Add the total of high and critical vulnerabilities
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("\n")
                            
                            high_severity_count = cve_infos.get('high_severity_count', 0)
                            critical_severity_count = cve_infos.get('critical_severity_count', 0)
                            
                            if high_severity_count > 0 or critical_severity_count > 0:
                                run = p.add_run("Résumé :\n")
                                run.bold = True

                                run = p.add_run(f"- ")
                                run = p.add_run(f"{high_severity_count} ")
                                run.bold = True
                                run = p.add_run(f' vulnérabilités ')
                                run = p.add_run(f'ÉLEVÉE \n')
                                run.italic = True

                                run = p.add_run(f"- ")
                                run = p.add_run(f"{critical_severity_count} ")
                                run.bold = True
                                run = p.add_run(f' vulnérabilités ')
                                run = p.add_run(f'CRITIQUE \n')
                                run.italic = True
                            
                            # Add a paragraph with the most impactful vulnerability
                            most_impactful_vuln = cve_infos.get('most_impactful_vuln', None)
                            if most_impactful_vuln is not None:
                                p = d.add_paragraph()
                                p.style = d.styles['Normal']
                                run = p.add_run("Description de la vulnérabilité ayant le plus d'impact (")
                                run.bold = True
                                hyperlink = add_hyperlink(p, CVE_URL + most_impactful_vuln['id'], most_impactful_vuln['id'], '0000FF', True)
                                run = p.add_run(") :\n")
                                run.bold = True
                                run = p.add_run('"')
                                run = p.add_run(most_impactful_vuln['description'])
                                run.italic = True
                                run = p.add_run('"')
                        else:
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Aucune vulnérabilité n'a été trouvée")
                            
                        time.sleep(2.5) # Sleep to avoid rate limit on CVE API
                        
                    except Exception as e:
                        print(f"[!] Error while writing CVE : {e}")
                        p = d.add_paragraph()
                        p.style = d.styles['Normal']
                        run = p.add_run("Aucune vulnérabilité n'a été trouvée")
                        # run.bold = True
                    
                    
                ##########################
                ###### FIN CUSTOM ########
                ##########################    
                        
                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")

                d.add_page_break()
            else:
                # This target had no clicked or submitted events so move on to next
                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
                continue

        print("[+] Finished writing Detailed Analysis section...")
        # End of the detailed results and the beginning of browser, location, and OS stats
        d.add_heading("Statistics", 1)
        p = d.add_paragraph("Le tableau suivant présente les navigateurs observés :")
        # Create browser table
        browser_table = d.add_table(rows=1, cols=4, style="GoReport")
        self._set_word_column_width(browser_table.columns[0], Cm(6.24))
        self._set_word_column_width(browser_table.columns[1], Cm(2.0))
        self._set_word_column_width(browser_table.columns[2], Cm(2.0))
        self._set_word_column_width(browser_table.columns[3], Cm(7.7))

        header1 = browser_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Navigateur", "Cell Text").bold = True

        header2 = browser_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Vu", "Cell Text").bold = True
        
        if WRITE_CVE_PARTS:
            header3 = browser_table.cell(0, 2)
            header3.text = ""
            header3.paragraphs[0].add_run("À jour", "Cell Text").bold = True
            
            header4 = browser_table.cell(0, 3)
            header4.text = ""
            header4.paragraphs[0].add_run("Vulnérabilités (High / Critical)", "Cell Text").bold = True

        p = d.add_paragraph("\nLe tableau suivant présente les systèmes d'exploitation observés :")

        # Create OS table
        os_table = d.add_table(rows=1, cols=3, style="GoReport")
        # self._set_word_column_width(os_table.columns[0], Cm(6.24))
        # self._set_word_column_width(os_table.columns[1], Cm(2.0))
        # self._set_word_column_width(os_table.columns[2], Cm(2.0))
        # self._set_word_column_width(os_table.columns[3], Cm(7.7))
        os_table.autofit = True
        os_table.allow_autofit = True

        header1 = os_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("OS", "Cell Text").bold = True

        header2 = os_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Vu", "Cell Text").bold = True
        
        if WRITE_CVE_PARTS:
            header3 = os_table.cell(0, 2)
            header3.text = ""
            header3.paragraphs[0].add_run("À jour", "Cell Text").bold = True

        p = d.add_paragraph("\nLe tableau suivant présente les localisations observées :")

        # Create geo IP table
        location_table = d.add_table(rows=1, cols=2, style="GoReport")
        self._set_word_column_width(location_table.columns[0], Cm(7.24))
        self._set_word_column_width(location_table.columns[1], Cm(3.35))

        header1 = location_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Localisation", "Cell Text").bold = True

        header2 = location_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Visites", "Cell Text").bold = True

        p = d.add_paragraph("\nLe tableau suivant présente les adresses IP capturées :")

        # Create IP address table
        ip_add_table = d.add_table(rows=1, cols=2, style="GoReport")
        self._set_word_column_width(ip_add_table.columns[0], Cm(7.24))
        self._set_word_column_width(ip_add_table.columns[1], Cm(3.35))

        header1 = ip_add_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Adresse IP", "Cell Text").bold = True

        header2 = ip_add_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Vu", "Cell Text").bold = True

        p = d.add_paragraph("\nLe tableau suivant présente les adresses IP correspondant aux données de géolocalisation :")

        # Create IP address and location table
        ip_loc_table = d.add_table(rows=1, cols=2, style="GoReport")
        self._set_word_column_width(ip_loc_table.columns[0], Cm(7.24))
        self._set_word_column_width(ip_loc_table.columns[1], Cm(3.35))

        header1 = ip_loc_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Adresse IP", "Cell Text").bold = True

        header2 = ip_loc_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Localisation", "Cell Text").bold = True

        # Counters are used here again to track rows
        counter = 1
        # Counter is used to count all elements in the lists to create a unique list with totals
        counted_browsers = Counter(self.browsers)
        for key, value in counted_browsers.items():
            browser_table.add_row()
            cell = browser_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = browser_table.cell(counter, 1)
            cell.text = f"{value}"
            
            if WRITE_CVE_PARTS:
                try:
                    # Check if the browser is up-to-date and add a checkmark or cross
                    cell = browser_table.cell(counter, 2)
                    if sofware_infos_dict.get(key, {}).get('version_browser_outdated', False):
                        cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")
                    else:
                        cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
                    
                    # Check if the browser has vulnerabilities and add a checkmark or cross
                    cell = browser_table.cell(counter, 3)
                    high_severity_count = sofware_infos_dict.get(key, {}).get('high_severity_count', 0)
                    critical_severity_count = sofware_infos_dict.get(key, {}).get('critical_severity_count', 0)
                    total_vulnerabilities = high_severity_count + critical_severity_count
                    if total_vulnerabilities > 0:
                        cell.paragraphs[0].add_run(str(high_severity_count), "Cell Text Miss")
                        cell.paragraphs[0].add_run(" / ").font.color.rgb = RGBColor(0x70, 0x70, 0x70)
                        cell.paragraphs[0].add_run(str(critical_severity_count)).bold = True
                    else:
                        cell.paragraphs[0].add_run("0", "Cell Text Hit")
                except Exception as e:
                    pass
            
            counter += 1

        counter = 1
        counted_os = Counter(self.operating_systems)
        for key, value in counted_os.items():
            os_table.add_row()
            cell = os_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = os_table.cell(counter, 1)
            cell.text = f"{value}"
            
            if WRITE_CVE_PARTS:
                try:
                    # Check if the OS is up-to-date and add a checkmark or cross
                    cell = os_table.cell(counter, 2)
                    if sofware_infos_dict.get(key, {}).get('version_os_outdated', False):
                        cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")
                    else:
                        cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
                except Exception as e:
                    pass
            
            counter += 1

        counter = 1
        counted_locations = Counter(self.locations)
        for key, value in counted_locations.items():
            location_table.add_row()
            cell = location_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = location_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_ip_addresses = Counter(self.ip_addresses)
        for key, value in counted_ip_addresses.items():
            ip_add_table.add_row()
            cell = ip_add_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = ip_add_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        for key, value in self.ip_and_location.items():
            ip_loc_table.add_row()
            cell = ip_loc_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = ip_loc_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        # Finalize document and save it as the value of output_word_report
        d.save(f"{self.output_word_report}")
        print(f"[+] Done! Check \"{self.output_word_report}\" for your results.")

    def config_section_map(self, config_parser, section):
        """This function helps by reading accepting a config file section, from gophish.config,
        and returning a dictionary object that can be referenced for configuration settings.
        """
        section_dict = {}
        options = config_parser.options(section)
        for option in options:
            try:
                section_dict[option] = config_parser.get(section, option)
                if section_dict[option] == -1:
                    print(f"[-] Skipping: {option}")
            except:
                print(f"[!] There was an error with: {option}")
                section_dict[option] = None
        return section_dict
