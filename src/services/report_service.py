import subprocess
import os
import sys
import re
from flask import send_file
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt
from io import BytesIO
import ollama
from src.config import Config
from src.lib.goreport_lib import Goreport
from src.services.llm_service import generate_ai_analysis

def get_latest_goreport_docx(campaign_id):
    reports_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "reports"))
    candidates = sorted(Path(reports_dir).glob(f"rapport_campagne_{campaign_id}_*.docx"), key=os.path.getmtime, reverse=True)
    return candidates[0] if candidates else None

def generate_docx_with_goreport(campaign_id, force=False):
    print(f" Génération GoReport pour la campagne {campaign_id}")

    # Dossier de stockage des rapports
    reports_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "reports"))
    reports_path = Path(reports_dir)
    pattern = f"rapport_campagne_{campaign_id}_*.docx"

    # Si force est True, supprimer les anciens rapports pour cette campagne
    if force:
        for f in reports_path.glob(pattern):
            try:
                os.remove(f)
                print(f" Suppression de l'ancien rapport : {f}")
            except Exception as e:
                print("Erreur lors de la suppression :", e)

    # Lancement de GoReport pour générer le rapport de base
  

    result = subprocess.run(
        [sys.executable, "-m", "src.services.goreport_service", "--id", str(campaign_id), "--format", "word"],
        cwd=os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")),
        capture_output=True,
        text=True
    )

    if result.returncode != 0:
        print(" GoReport Error:", result.stderr)
        raise RuntimeError("Erreur GoReport : " + result.stderr)

    print(" GoReport exécuté avec succès")

    # Récupérer le dernier fichier généré
    latest_file = None
    for f in sorted(reports_path.glob(pattern), key=os.path.getmtime, reverse=True):
        latest_file = f
        break

    if not latest_file:
        raise FileNotFoundError("Aucun fichier DOCX généré par GoReport")

    print(f"Rapport trouvé : {latest_file}")

    # Génération du texte d'analyse via l'IA
    ai_data = generate_ai_analysis(campaign_id)
    # Ouvrir le document généré pour y ajouter le rapport d'analyse
    document = Document(str(latest_file))
        # Texte d’analyse via IA


    document.add_page_break()
    document.add_heading("Analyse de la campagne", level=1)

    document.add_heading("Analyse des résultats généraux", level=2)
    document.add_paragraph(ai_data.analyse)

    document.add_heading("Recommandations", level=2)
    for i, reco in enumerate(ai_data.recommandations, 1):
        document.add_paragraph(f"{i}. {reco.titre.upper()}\n   {reco.explication}")

    document.add_heading("Conclusion", level=2)
    document.add_paragraph(ai_data.conclusion)

    # Sauvegarder le document final
    document.save(str(latest_file))
    
    # Retourner le fichier final
    return send_file(
        str(latest_file),
        as_attachment=True,
        download_name=latest_file.name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def download_report_styled(campaign_id):
    """
    Cherche le dernier rapport Word généré pour une campagne donnée et le renvoie en téléchargement.
    """
    # Répertoire contenant les rapports générés
    reports_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "reports"))

    # Récupérer les rapports triés par date (du plus récent au plus ancien)
    candidates = sorted(
        Path(reports_dir).glob(f"rapport_campagne_{campaign_id}_*.docx"),
        key=os.path.getmtime,
        reverse=True
    )

    latest_file = candidates[0] if candidates else None
    if not latest_file:
        return "Rapport introuvable", 404

    # Envoyer le fichier en téléchargement
    return send_file(
        str(latest_file),
        as_attachment=True,
        download_name=latest_file.name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )